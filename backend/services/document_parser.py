"""Document parsing and ChunkNorris-style structure-aware chunking."""

from __future__ import annotations

import logging
import os
import re
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# ChunkNorris-style contract version — bump when chunk boundaries change.
CHUNK_VERSION = "cn_v1"
SOFT_CHUNK_CHARS = 1800
HARD_CHUNK_CHARS = 2400
MIN_USEFUL_CHARS = 40

_NUMBERED_HEADING_RE = re.compile(r"^\d+(\.\d+)*\s+[А-ЯA-ZЁ]")
_CODE_FENCE_RE = re.compile(r"^\s*```")
_CODE_KEYWORD_RE = re.compile(
    r"(?:^|\n)\s*(?:import\s+\w+|from\s+\w+\s+import|def\s+\w+\s*\(|class\s+\w+\s*[:\(]|"
    r"function\s+\w+\s*\(|console\.log|#!/usr/bin|package\s+\w+)",
    re.IGNORECASE | re.MULTILINE,
)
_CODE_PUNCT_RE = re.compile(r"[{};=<>]|::|->|=>")


def slugify_filename(filename: str, max_len: int = 32) -> str:
    """Build a stable uppercase slug from a file basename (no extension)."""
    base = os.path.splitext(os.path.basename(filename))[0]
    slug = re.sub(r"[^A-Za-z0-9\u0400-\u04FF]+", "-", base.upper())
    slug = re.sub(r"-+", "-", slug).strip("-")
    if not slug:
        slug = "DOC"
    return slug[:max_len].rstrip("-")


def normalize_code(code_str: str) -> str:
    """Normalizes document codes, padding the middle index to 2 digits if necessary."""
    clean = re.sub(r"[\s\-\–\—\u2013\u2014\x96]+", "-", code_str)
    match = re.match(r"^([A-ZА-Яa-zа-я]+)-(\d+)-(\d{4})$", clean)
    if match:
        prefix, num, year = match.groups()
        if len(num) == 1:
            num = f"0{num}"
        return f"{prefix.upper()}-{num}-{year}"
    return clean.upper()


def is_code_like(text: str, *, monospace: bool = False) -> bool:
    """Conservative heuristic: skip source-code fragments, keep formulas when unsure."""
    stripped = (text or "").strip()
    if not stripped:
        return False
    if _CODE_FENCE_RE.search(stripped):
        return True
    lines = [ln for ln in stripped.splitlines() if ln.strip()]
    keyword_hits = len(_CODE_KEYWORD_RE.findall(stripped))
    if monospace and keyword_hits >= 1:
        return True
    if keyword_hits >= 2 and len(lines) >= 2:
        return True
    if keyword_hits >= 1 and _CODE_PUNCT_RE.search(stripped):
        punct_ratio = len(_CODE_PUNCT_RE.findall(stripped)) / max(len(stripped), 1)
        if punct_ratio >= 0.015 or len(lines) >= 2:
            return True
    return False


def _heading_level_from_style(style_name: str) -> Optional[int]:
    if not style_name:
        return None
    if style_name.startswith("Heading"):
        m = re.search(r"(\d+)", style_name)
        if m:
            return max(1, min(int(m.group(1)), 6))
        return 1
    return None


def _is_numbered_heading(text: str) -> bool:
    return len(text) < 120 and bool(_NUMBERED_HEADING_RE.match(text))


def format_table_rows(rows: Sequence[Sequence[str]]) -> str:
    """Serialize table rows as pipe-delimited lines."""
    lines: List[str] = []
    for row in rows:
        cells = [re.sub(r"\s+", " ", (c or "").strip()) for c in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def split_table_with_repeated_headers(
    header_rows: Sequence[Sequence[str]],
    body_rows: Sequence[Sequence[str]],
    *,
    soft_limit: int = SOFT_CHUNK_CHARS,
    hard_limit: int = HARD_CHUNK_CHARS,
) -> List[Tuple[str, int]]:
    """
    Split an oversized table into parts. Each part repeats header rows.

    Returns list of (serialized_text, part_index_1based).
    """
    header_text = format_table_rows(header_rows) if header_rows else ""
    header_len = len(header_text) + (1 if header_text else 0)

    if not body_rows:
        text = header_text
        return [(text, 1)] if text.strip() else []

    full = format_table_rows(list(header_rows) + list(body_rows)) if header_rows else format_table_rows(body_rows)
    if len(full) <= soft_limit:
        return [(full, 1)]

    parts: List[Tuple[str, int]] = []
    current: List[Sequence[str]] = []
    current_len = header_len
    part_idx = 1

    for row in body_rows:
        row_text = format_table_rows([row])
        row_cost = len(row_text) + 1
        if current and current_len + row_cost > soft_limit:
            body = format_table_rows(current)
            text = f"{header_text}\n{body}" if header_text else body
            parts.append((text, part_idx))
            part_idx += 1
            current = []
            current_len = header_len
        # Single row larger than hard limit: still emit alone with headers.
        if not current and header_len + row_cost > hard_limit:
            body = format_table_rows([row])
            text = f"{header_text}\n{body}" if header_text else body
            parts.append((text, part_idx))
            part_idx += 1
            continue
        current.append(row)
        current_len += row_cost

    if current:
        body = format_table_rows(current)
        text = f"{header_text}\n{body}" if header_text else body
        parts.append((text, part_idx))

    return parts


def _section_label(section_path: Sequence[str]) -> str:
    if not section_path:
        return "Введение"
    return section_path[-1]


def _prepend_section_context(text: str, section_path: Sequence[str]) -> str:
    if not section_path:
        return text
    header_block = "\n".join(section_path)
    if text.startswith(header_block):
        return text
    return f"{header_block}\n\n{text}"


def _make_chunk(
    *,
    index: int,
    text: str,
    section_path: Sequence[str],
    page: Optional[int],
    content_type: str,
    source_block_id: str,
    table_header: Optional[str] = None,
    table_part: Optional[int] = None,
) -> Dict[str, Any]:
    return {
        "index": index,
        "text": text,
        "section": _section_label(section_path),
        "section_path": list(section_path),
        "page": page,
        "content_type": content_type,
        "chunk_version": CHUNK_VERSION,
        "source_block_id": source_block_id,
        "table_header": table_header,
        "table_part": table_part,
    }


def assemble_chunks_from_blocks(
    blocks: Sequence[Dict[str, Any]],
    *,
    soft_limit: int = SOFT_CHUNK_CHARS,
    hard_limit: int = HARD_CHUNK_CHARS,
) -> List[Dict[str, Any]]:
    """
    Convert normalized blocks into versioned chunks.

    Block keys:
      type: heading | text | table | code
      text / rows / level / page / block_id
    """
    chunks: List[Dict[str, Any]] = []
    section_stack: List[Tuple[int, str]] = []
    buffer: List[str] = []
    buffer_len = 0
    buffer_page: Optional[int] = None
    buffer_block_ids: List[str] = []
    skipped_code = 0

    def current_path() -> List[str]:
        return [title for _, title in section_stack]

    def flush_text_buffer() -> None:
        nonlocal buffer, buffer_len, buffer_page, buffer_block_ids
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        buffer = []
        buffer_len = 0
        page = buffer_page
        buffer_page = None
        block_id = buffer_block_ids[0] if buffer_block_ids else f"text-{len(chunks)}"
        buffer_block_ids = []
        if len(text) < MIN_USEFUL_CHARS:
            return
        path = current_path()
        text_with_ctx = _prepend_section_context(text, path)
        # Soft-split oversized prose by paragraphs.
        if len(text_with_ctx) <= soft_limit:
            chunks.append(
                _make_chunk(
                    index=len(chunks),
                    text=text_with_ctx,
                    section_path=path,
                    page=page,
                    content_type="text",
                    source_block_id=block_id,
                )
            )
            return
        paras = [p for p in text.split("\n") if p.strip()]
        cur: List[str] = []
        cur_len = 0
        part = 0
        ctx_prefix = "\n".join(path) + "\n\n" if path else ""
        ctx_len = len(ctx_prefix)
        for para in paras:
            cost = len(para) + 1
            if cur and cur_len + ctx_len + cost > soft_limit:
                body = "\n".join(cur)
                chunks.append(
                    _make_chunk(
                        index=len(chunks),
                        text=ctx_prefix + body,
                        section_path=path,
                        page=page,
                        content_type="text",
                        source_block_id=f"{block_id}-p{part}",
                    )
                )
                part += 1
                cur = []
                cur_len = 0
            if not cur and ctx_len + cost > hard_limit:
                chunks.append(
                    _make_chunk(
                        index=len(chunks),
                        text=ctx_prefix + para,
                        section_path=path,
                        page=page,
                        content_type="text",
                        source_block_id=f"{block_id}-p{part}",
                    )
                )
                part += 1
                continue
            cur.append(para)
            cur_len += cost
        if cur:
            body = "\n".join(cur)
            chunks.append(
                _make_chunk(
                    index=len(chunks),
                    text=ctx_prefix + body,
                    section_path=path,
                    page=page,
                    content_type="text",
                    source_block_id=f"{block_id}-p{part}" if part else block_id,
                )
            )

    for block in blocks:
        btype = block.get("type")
        if btype == "code":
            skipped_code += 1
            continue
        if btype == "heading":
            flush_text_buffer()
            level = int(block.get("level") or 1)
            title = (block.get("text") or "").strip()
            if not title:
                continue
            while section_stack and section_stack[-1][0] >= level:
                section_stack.pop()
            section_stack.append((level, title))
            continue
        if btype == "table":
            flush_text_buffer()
            rows = block.get("rows") or []
            if not rows:
                continue
            header_rows = rows[:1]
            body_rows = rows[1:] if len(rows) > 1 else []
            # If only one row, treat it as body with empty header.
            if not body_rows and header_rows:
                body_rows = list(header_rows)
                header_rows = []
            header_serialized = format_table_rows(header_rows) if header_rows else None
            parts = split_table_with_repeated_headers(
                header_rows,
                body_rows,
                soft_limit=soft_limit,
                hard_limit=hard_limit,
            )
            path = current_path()
            base_id = block.get("block_id") or f"table-{len(chunks)}"
            for text, part_idx in parts:
                text_with_ctx = _prepend_section_context(text, path)
                chunks.append(
                    _make_chunk(
                        index=len(chunks),
                        text=text_with_ctx,
                        section_path=path,
                        page=block.get("page"),
                        content_type="table",
                        source_block_id=f"{base_id}-part{part_idx}" if len(parts) > 1 else base_id,
                        table_header=header_serialized,
                        table_part=part_idx if len(parts) > 1 else None,
                    )
                )
            continue
        if btype == "text":
            text = (block.get("text") or "").strip()
            if not text:
                continue
            if is_code_like(text, monospace=bool(block.get("monospace"))):
                skipped_code += 1
                continue
            if buffer and buffer_len + len(text) + 1 > soft_limit:
                flush_text_buffer()
            buffer.append(text)
            buffer_len += len(text) + 1
            if buffer_page is None:
                buffer_page = block.get("page")
            bid = block.get("block_id")
            if bid:
                buffer_block_ids.append(str(bid))
            continue

    flush_text_buffer()
    if skipped_code:
        logger.info("Skipped %d code-like block(s) during chunk assembly", skipped_code)
    return chunks


def _iter_docx_blocks(doc: Document) -> Iterable[Any]:
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def _docx_table_rows(table: DocxTable) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _docx_paragraph_is_code(paragraph: DocxParagraph) -> bool:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    if "Code" in style_name or style_name in {"HTML Preformatted", "Source Code"}:
        return True
    text = paragraph.text or ""
    monospace = False
    for run in paragraph.runs:
        font_name = (run.font.name or "").lower() if run.font else ""
        if any(m in font_name for m in ("consolas", "courier", "mono", "menlo", "monaco")):
            monospace = True
            break
    return is_code_like(text, monospace=monospace)


class DocumentParser:
    def __init__(self, target_categories: List[str] = ["Обзоры", "Статьи", "Доклады"]):
        self.target_categories = target_categories

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parses a DOCX file, extracts metadata and structure-aware chunks."""
        doc = Document(file_path)

        title = ""
        code = ""
        year = None
        authors: List[str] = []
        approver = ""

        filename = os.path.basename(file_path)
        year_match = re.search(r"\b(20\d{2})\b", filename)
        if year_match:
            year = int(year_match.group(1))

        code_regex = r"\b([ОИП|ТИ|ИС|НДС]+[\s\-\–\—\u2013\u2014\x96]*?\d+[\s\-\–\—\u2013\u2014\x96]*?\d{4})\b"
        code_match = re.search(code_regex, filename, re.IGNORECASE)
        if code_match:
            code = normalize_code(code_match.group(1))

        meta_paras: List[str] = []
        for i, p in enumerate(doc.paragraphs[:60]):
            text = p.text.strip()
            if not text:
                continue
            meta_paras.append(text)

            if not code:
                code_p_match = re.search(code_regex, text, re.IGNORECASE)
                if code_p_match:
                    code = normalize_code(code_p_match.group(1))

            if "УТВЕРЖДАЮ" in text or "Утверждаю" in text:
                for next_p in doc.paragraphs[i + 1 : i + 5]:
                    nt = next_p.text.strip()
                    if nt and any(
                        role in nt
                        for role in ["Директор", "Начальник", "Заместитель", "к.г.-м.н.", "Цымбулов", "Козырев"]
                    ):
                        approver = nt
                        break

            if "исполнителей" in text.lower() or "разработчик" in text.lower():
                for next_p in doc.paragraphs[i + 1 : i + 10]:
                    nt = next_p.text.strip()
                    if nt and any(
                        kw in nt.lower()
                        for kw in ["специалист", "инженер", "научный сотрудник", "разработал", "выполнил"]
                    ):
                        name_match = re.search(r"([А-ЯA-Z][а-яa-z]+\s+[А-ЯA-Z]\.[А-ЯA-Z]\.)", nt)
                        if name_match:
                            authors.append(name_match.group(1))

        for p in meta_paras[:15]:
            p_clean = p.strip()
            if len(p_clean) > 12 and not re.search(r"_{3,}|-{3,}|–{3,}|—{3,}", p_clean):
                if not any(
                    x in p_clean.lower()
                    for x in [
                        "утверждаю",
                        "директор",
                        "институт",
                        "департамент",
                        "список",
                        "обзор",
                        "оип",
                        "ти-",
                        "ис-",
                        "оглавление",
                    ]
                ):
                    title = p_clean
                    break

        if not title:
            title = os.path.splitext(filename)[0]

        blocks: List[Dict[str, Any]] = []
        block_i = 0
        for item in _iter_docx_blocks(doc):
            if isinstance(item, DocxParagraph):
                text = item.text.strip()
                if not text:
                    continue
                if _docx_paragraph_is_code(item):
                    blocks.append({"type": "code", "text": text, "block_id": f"docx-code-{block_i}"})
                    block_i += 1
                    continue
                style_name = (item.style.name or "") if item.style else ""
                level = _heading_level_from_style(style_name)
                if level is None and _is_numbered_heading(text):
                    dots = text.split()[0].count(".")
                    level = min(dots + 1, 6)
                if level is not None:
                    blocks.append(
                        {
                            "type": "heading",
                            "text": text,
                            "level": level,
                            "block_id": f"docx-h-{block_i}",
                        }
                    )
                else:
                    blocks.append(
                        {
                            "type": "text",
                            "text": text,
                            "block_id": f"docx-t-{block_i}",
                        }
                    )
                block_i += 1
            elif isinstance(item, DocxTable):
                rows = _docx_table_rows(item)
                if rows:
                    blocks.append(
                        {
                            "type": "table",
                            "rows": rows,
                            "block_id": f"docx-tbl-{block_i}",
                        }
                    )
                    block_i += 1

        chunks = assemble_chunks_from_blocks(blocks)

        if not year:
            for chunk in chunks[:3]:
                year_match = re.search(r"\b(20\d{2})\b", chunk["text"])
                if year_match:
                    year = int(year_match.group(1))
                    break
        if not year:
            year = 2024

        return {
            "file_path": file_path,
            "filename": filename,
            "file_slug": slugify_filename(filename),
            "title": title,
            "code": code or "N/A",
            "year": year,
            "authors": list(set(authors)) if authors else ["Не указан"],
            "approver": approver or "Не указан",
            "format": "docx",
            "chunk_version": CHUNK_VERSION,
            "chunks": chunks,
        }

    def _pdf_extract_blocks(self, doc: fitz.Document) -> List[Dict[str, Any]]:
        """Layout-light PDF → heading/text/table/code blocks."""
        blocks: List[Dict[str, Any]] = []
        block_i = 0

        # Body font size ≈ median of span sizes across first pages.
        sizes: List[float] = []
        for page_num in range(min(doc.page_count, 5)):
            page = doc[page_num]
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        sz = float(span.get("size") or 0)
                        if sz > 0:
                            sizes.append(sz)
        body_size = sorted(sizes)[len(sizes) // 2] if sizes else 11.0

        for page_num in range(doc.page_count):
            page = doc[page_num]
            page_index = page_num + 1

            # Prefer native table finder when available.
            table_rects: List[fitz.Rect] = []
            try:
                finder = page.find_tables()
                tables = list(finder.tables) if finder is not None else []
            except Exception:
                tables = []

            for ti, table in enumerate(tables):
                try:
                    extracted = table.extract()
                except Exception:
                    continue
                rows = []
                for row in extracted or []:
                    cells = [("" if c is None else str(c)).strip() for c in row]
                    if any(cells):
                        rows.append(cells)
                if not rows:
                    continue
                try:
                    table_rects.append(fitz.Rect(table.bbox))
                except Exception:
                    pass
                blocks.append(
                    {
                        "type": "table",
                        "rows": rows,
                        "page": page_index,
                        "block_id": f"pdf-tbl-{page_index}-{ti}",
                    }
                )

            for bi, block in enumerate(page.get_text("dict").get("blocks", [])):
                if block.get("type") != 0:
                    continue
                bbox = fitz.Rect(block.get("bbox", (0, 0, 0, 0)))
                if any(bbox.intersects(tr) for tr in table_rects):
                    continue

                line_texts: List[str] = []
                max_size = 0.0
                mono_hits = 0
                span_count = 0
                for line in block.get("lines", []):
                    parts: List[str] = []
                    for span in line.get("spans", []):
                        span_count += 1
                        txt = (span.get("text") or "").strip()
                        if not txt:
                            continue
                        parts.append(txt)
                        max_size = max(max_size, float(span.get("size") or 0))
                        font = (span.get("font") or "").lower()
                        if any(m in font for m in ("mono", "consolas", "courier", "menlo")):
                            mono_hits += 1
                    if parts:
                        line_texts.append(" ".join(parts))
                text = "\n".join(line_texts).strip()
                if not text:
                    continue

                monospace = span_count > 0 and mono_hits / span_count >= 0.5
                if is_code_like(text, monospace=monospace):
                    blocks.append(
                        {
                            "type": "code",
                            "text": text,
                            "page": page_index,
                            "block_id": f"pdf-code-{page_index}-{bi}",
                        }
                    )
                    continue

                is_heading = False
                level = 1
                if len(text) < 140 and "\n" not in text:
                    if max_size >= body_size + 1.5:
                        is_heading = True
                        level = 1 if max_size >= body_size + 3 else 2
                    elif _is_numbered_heading(text):
                        is_heading = True
                        dots = text.split()[0].count(".")
                        level = min(dots + 1, 6)

                if is_heading:
                    blocks.append(
                        {
                            "type": "heading",
                            "text": text,
                            "level": level,
                            "page": page_index,
                            "block_id": f"pdf-h-{page_index}-{bi}",
                        }
                    )
                else:
                    blocks.append(
                        {
                            "type": "text",
                            "text": text,
                            "page": page_index,
                            "monospace": monospace,
                            "block_id": f"pdf-t-{page_index}-{bi}",
                        }
                    )
                block_i += 1

        return blocks

    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parses a PDF file with layout-light header/table detection and semantic chunking."""
        doc = fitz.open(file_path)
        filename = os.path.basename(file_path)

        title = os.path.splitext(filename)[0]
        code = ""
        year = None

        year_match = re.search(r"\b(20\d{2})\b", filename)
        if year_match:
            year = int(year_match.group(1))

        code_regex = r"\b([ОИП|ТИ|ИС|НДС]+[\s\-\–\—\u2013\u2014\x96]*?\d+[\s\-\–\—\u2013\u2014\x96]*?\d{4})\b"
        code_match = re.search(code_regex, filename, re.IGNORECASE)
        if code_match:
            code = normalize_code(code_match.group(1))

        if doc.page_count > 0:
            first_text = doc[0].get_text()
            if not code:
                code_p_match = re.search(code_regex, first_text, re.IGNORECASE)
                if code_p_match:
                    code = normalize_code(code_p_match.group(1))
            if not year:
                year_match = re.search(r"\b(20\d{2})\b", first_text)
                if year_match:
                    year = int(year_match.group(1))

        blocks = self._pdf_extract_blocks(doc)
        chunks = assemble_chunks_from_blocks(blocks)

        if not year:
            year = 2024

        return {
            "file_path": file_path,
            "filename": filename,
            "file_slug": slugify_filename(filename),
            "title": title,
            "code": code or "N/A",
            "year": year,
            "authors": ["Не указан"],
            "approver": "Не указан",
            "format": "pdf",
            "chunk_version": CHUNK_VERSION,
            "chunks": chunks,
        }

    def parse_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Determines format and parses the file."""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".docx":
                return self.parse_docx(file_path)
            if ext == ".pdf":
                return self.parse_pdf(file_path)
            return None
        except Exception as e:
            logger.warning("Error parsing file %s: %s", file_path, e)
            return None

    def scan_directory(self, base_dir: str) -> List[str]:
        """Scans directory and returns list of supported files to parse."""
        supported_files = []
        for root, _dirs, files in os.walk(base_dir):
            category_match = False
            for cat in self.target_categories:
                if cat in root:
                    category_match = True
                    break

            if not category_match and "Источники информации" in root:
                category_match = True

            if category_match:
                for file in files:
                    ext = os.path.splitext(file)[1].lower()
                    if ext in [".docx", ".pdf"]:
                        supported_files.append(os.path.join(root, file))
        return supported_files
